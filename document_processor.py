import os
import re
from typing import Optional
import PyPDF2
import docx
from io import BytesIO

class DocumentProcessor:
    def __init__(self):
        """Initialize the document processor."""
        pass
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or Word documents."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_text_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
        except Exception as e:
            # Fallback to alternative PDF processing if PyPDF2 fails
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                raise Exception(f"Failed to extract text from PDF: {str(e)}. Consider installing pdfplumber for better PDF support.")
            except Exception as fallback_error:
                raise Exception(f"Failed to extract text from PDF with both methods: {str(e)} and {str(fallback_error)}")
        
        return text.strip()
    
    def _extract_text_from_word(self, file_path: str) -> str:
        """Extract text from Word document."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def extract_title(self, content: str) -> str:
        """Extract title from document content."""
        if not content.strip():
            return "Untitled Document"
        
        lines = content.split('\n')
        
        # Strategy 1: Look for the first non-empty line
        for line in lines:
            line = line.strip()
            if line and len(line) > 3:  # Minimum title length
                # Clean up the title
                title = self._clean_title(line)
                if title and len(title) <= 200:  # Reasonable title length
                    return title
        
        # Strategy 2: Look for lines that might be titles (all caps, short, etc.)
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            if line and self._is_likely_title(line):
                title = self._clean_title(line)
                if title and len(title) <= 200:
                    return title
        
        # Strategy 3: Use first sentence if no clear title found
        sentences = content.split('.')
        if sentences:
            first_sentence = sentences[0].strip()
            if first_sentence and len(first_sentence) <= 200:
                return self._clean_title(first_sentence)
        
        # Fallback: Use first 100 characters
        fallback_title = content[:100].strip()
        return self._clean_title(fallback_title) if fallback_title else "Untitled Document"
    
    def _clean_title(self, title: str) -> str:
        """Clean and format title text."""
        # Remove extra whitespace
        title = ' '.join(title.split())
        
        # Remove common file artifacts
        title = re.sub(r'^(Title:|TITLE:|Subject:|SUBJECT:)', '', title, flags=re.IGNORECASE)
        
        # Remove special characters at the beginning and end
        title = title.strip(' .-_=+')
        
        # Capitalize properly if all caps or all lowercase
        if title.isupper() or title.islower():
            title = title.title()
        
        return title
    
    def _is_likely_title(self, line: str) -> bool:
        """Determine if a line is likely to be a title."""
        line = line.strip()
        
        # Check for title characteristics
        if len(line) < 5 or len(line) > 100:
            return False
        
        # Likely title if:
        # - All caps
        # - Starts with capital and has few lowercase letters
        # - Contains common title words
        
        if line.isupper():
            return True
        
        title_indicators = ['chapter', 'section', 'introduction', 'conclusion', 'abstract', 'summary']
        if any(indicator in line.lower() for indicator in title_indicators):
            return True
        
        # Check if it's a short line with title-like formatting
        words = line.split()
        if len(words) <= 8 and line[0].isupper():
            capital_count = sum(1 for c in line if c.isupper())
            if capital_count / len(line) > 0.1:  # At least 10% capitals
                return True
        
        return False
    
    def extract_metadata(self, file_path: str) -> dict:
        """Extract metadata from document."""
        metadata = {
            'file_size': os.path.getsize(file_path),
            'file_extension': os.path.splitext(file_path)[1].lower(),
            'filename': os.path.basename(file_path)
        }
        
        file_extension = metadata['file_extension']
        
        if file_extension == '.pdf':
            metadata.update(self._extract_pdf_metadata(file_path))
        elif file_extension in ['.docx', '.doc']:
            metadata.update(self._extract_word_metadata(file_path))
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> dict:
        """Extract metadata from PDF file."""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['creator'] = pdf_reader.metadata.get('/Creator', '')
                    metadata['producer'] = pdf_reader.metadata.get('/Producer', '')
                    metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
                    metadata['pdf_title'] = pdf_reader.metadata.get('/Title', '')
                
                metadata['page_count'] = len(pdf_reader.pages)
                
        except Exception:
            # If metadata extraction fails, return empty metadata
            metadata['page_count'] = 0
        
        return metadata
    
    def _extract_word_metadata(self, file_path: str) -> dict:
        """Extract metadata from Word document."""
        metadata = {}
        try:
            doc = docx.Document(file_path)
            
            if doc.core_properties:
                metadata['author'] = doc.core_properties.author or ''
                metadata['title'] = doc.core_properties.title or ''
                metadata['subject'] = doc.core_properties.subject or ''
                metadata['created'] = str(doc.core_properties.created) if doc.core_properties.created else ''
                metadata['modified'] = str(doc.core_properties.modified) if doc.core_properties.modified else ''
            
            # Count paragraphs as a rough page estimate
            metadata['paragraph_count'] = len(doc.paragraphs)
            
        except Exception:
            # If metadata extraction fails, return empty metadata
            metadata['paragraph_count'] = 0
        
        return metadata
